#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""The structure sidecar, tested on a Word file this test builds.

Word keeps a clause's number as a property of the paragraph (`w:numPr`), not as text, so it
never reaches the markdown: a real nolikums of 59,499 characters came through with four
heading lines, three of them converter artefacts, and not one clause number — while its own
text cited "Nolikuma 5.6.punktu". Counting paragraphs to put the numbers back is measurably
wrong, because the nesting is flattened too and the counter runs ahead.

These tests hold what the sidecar promises: the text is untouched, the index lines up with
what actually reached the markdown, the numbering definitions travel so a consumer can compute
a number rather than guess one, and a document with no numbering gets no sidecar at all.
"""

import io
import json
import os
import shutil
import sys
import tempfile
import unittest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import normalize


def a_document(path):
    """A .docx with a numbered clause tree, an unnumbered paragraph, and an empty one."""
    import docx
    d = docx.Document()
    d.add_paragraph("NOLIKUMS")                       # index 0, no numbering
    d.add_paragraph("")                               # skipped entirely
    for text, ilvl in (("Iepirkuma priekšmets", 0),
                       ("Tehniskās prasības", 0),
                       ("Pirmā apakšnodaļa", 1),
                       ("Otrā apakšnodaļa", 1)):
        p = d.add_paragraph(text)
        _number(p, num_id="7", ilvl=ilvl)
    d.add_paragraph("Pielikumā: darbu apjomi")        # unnumbered again
    d.save(path)


def _number(paragraph, num_id, ilvl):
    """Give a paragraph the numbering properties Word would, at the XML level."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    for tag, value in (("w:ilvl", str(ilvl)), ("w:numId", str(num_id))):
        el = OxmlElement(tag)
        el.set(qn("w:val"), value)
        numPr.append(el)
    pPr.append(numPr)


def _style_numbered(document, name, num_id, ilvl, based_on=None):
    """A paragraph style that carries the numbering, the way a Latvian nolikums template does."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    style = document.styles.add_style(name, 1)
    if based_on is not None:
        based = OxmlElement("w:basedOn")
        based.set(qn("w:val"), based_on)
        style._element.append(based)
        return style
    pPr = OxmlElement("w:pPr")
    numPr = OxmlElement("w:numPr")
    for tag, value in (("w:ilvl", str(ilvl)), ("w:numId", str(num_id))):
        el = OxmlElement(tag)
        el.set(qn("w:val"), value)
        numPr.append(el)
    pPr.append(numPr)
    style._element.append(pPr)
    return style


class NumberingOnTheStyle(unittest.TestCase):
    """Word resolves a number from the paragraph, then its style, then what that is based on.

    A tender template routinely numbers through styles: almost no paragraph carries numbering
    of its own, while the text cites its own clauses by number and the styles carry the depth
    spelled into their names. Reading only the paragraph reports a document with no numbering
    at all, which is a confident wrong answer about a real clause tree.
    """

    def setUp(self):
        import docx
        self.root = tempfile.mkdtemp(prefix="eis_style_num_")
        d = docx.Document()
        parent = _style_numbered(d, "Nodala 1.1", num_id=9, ilvl=0)
        _style_numbered(d, "Nodala 1.1.1", None, None, based_on=parent.style_id)
        d.add_paragraph("Iepirkuma priekšmets", style="Nodala 1.1")
        d.add_paragraph("Apakšklauzula", style="Nodala 1.1.1")
        d.add_paragraph("Parasts teksts bez stila")
        path = os.path.join(self.root, "nolikums.docx")
        d.save(path)
        _md, _media, self.structure = normalize.docx_to_md(
            path, os.path.join(self.root, "media"), "n")

    def tearDown(self):
        shutil.rmtree(self.root, ignore_errors=True)

    def test_a_paragraph_inherits_its_number_from_its_style(self):
        first = self.structure["paragraphs"][0]
        self.assertEqual(first["numId"], "9")
        self.assertEqual(first["from"], "style:Nodala 1.1")

    def test_and_through_the_style_that_one_is_based_on(self):
        second = self.structure["paragraphs"][1]
        self.assertEqual(second["numId"], "9")
        self.assertEqual(second["style"], "Nodala 1.1.1")
        self.assertEqual(second["from"], "style:Nodala 1.1",
                         "the basedOn chain must name where the numbering actually came from")

    def test_a_paragraph_with_neither_is_still_absent(self):
        self.assertEqual(len(self.structure["paragraphs"]), 2)


class NumberingSwitchedOff(unittest.TestCase):
    """`numId=0` on the paragraph is Word's way of saying "not numbered", and it beats the style."""

    def test_an_explicit_zero_wins_over_the_style_and_is_reported_as_the_paragraphs_own(self):
        import docx
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        root = tempfile.mkdtemp(prefix="eis_num_off_")
        self.addCleanup(shutil.rmtree, root, True)
        d = docx.Document()
        _style_numbered(d, "Nodala 1.1", num_id=9, ilvl=0)
        p = d.add_paragraph("Numerācija noņemta", style="Nodala 1.1")
        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        el = OxmlElement("w:numId")
        el.set(qn("w:val"), "0")
        numPr.append(el)
        pPr.append(numPr)
        path = os.path.join(root, "off.docx")
        d.save(path)
        _md, _media, structure = normalize.docx_to_md(path, os.path.join(root, "m"), "o")
        entry = structure["paragraphs"][0]
        self.assertEqual(entry["numId"], "0")
        self.assertEqual(entry["from"], "paragraph")

    def test_a_missing_ilvl_is_null_and_never_a_filled_in_zero(self):
        # "the file does not say" and "top level" are different claims, and a consumer computing
        # a clause number from the second when the first is true gets a confident wrong answer.
        import docx
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        root = tempfile.mkdtemp(prefix="eis_no_ilvl_")
        self.addCleanup(shutil.rmtree, root, True)
        d = docx.Document()
        p = d.add_paragraph("Punkts bez līmeņa")
        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        el = OxmlElement("w:numId")
        el.set(qn("w:val"), "3")
        numPr.append(el)
        pPr.append(numPr)
        path = os.path.join(root, "noilvl.docx")
        d.save(path)
        _md, _media, structure = normalize.docx_to_md(path, os.path.join(root, "m"), "n")
        self.assertIsNone(structure["paragraphs"][0]["ilvl"])


class TheSidecar(unittest.TestCase):
    def setUp(self):
        self.root = tempfile.mkdtemp(prefix="eis_structure_")
        self.docx = os.path.join(self.root, "nolikums.docx")
        a_document(self.docx)
        self.md, self.media, self.structure = normalize.docx_to_md(
            self.docx, os.path.join(self.root, "media"), "nolikums")

    def tearDown(self):
        shutil.rmtree(self.root, ignore_errors=True)

    def test_the_text_is_what_it_always_was(self):
        # The sidecar exists precisely so that this line never has to change.
        self.assertIn("NOLIKUMS", self.md)
        self.assertIn("Pirmā apakšnodaļa", self.md)
        self.assertNotIn("structure", self.md)
        self.assertNotIn("numId", self.md)

    def test_it_records_the_numbered_paragraphs_and_only_those(self):
        self.assertIsNotNone(self.structure)
        paras = self.structure["paragraphs"]
        self.assertEqual([p["index"] for p in paras], [1, 2, 3, 4])
        self.assertEqual([p["ilvl"] for p in paras], [0, 0, 1, 1])
        self.assertEqual({p["numId"] for p in paras}, {"7"})

    def test_the_index_counts_what_reached_the_markdown(self):
        # The empty paragraph is skipped by the converter; an index that counted every
        # paragraph in the file would put the first clause at 2 and drift from there.
        lines = [line for line in self.md.split("\n\n") if line.strip()]
        for p in self.structure["paragraphs"]:
            self.assertLess(p["index"], len(lines))
        self.assertEqual(lines[self.structure["paragraphs"][0]["index"]],
                         "Iepirkuma priekšmets")

    def test_every_entry_carries_a_digest_of_its_own_line(self):
        import hashlib
        lines = [line for line in self.md.split("\n\n") if line.strip()]
        for p in self.structure["paragraphs"]:
            want = hashlib.sha1(lines[p["index"]].encode("utf-8")).hexdigest()[:8]
            self.assertEqual(p["digest"], want,
                             "a consumer must be able to prove the alignment, not assume it")

    def test_the_numbering_definitions_travel(self):
        # Depth and order alone cannot print a number: a list may start at 3 or restart. The
        # definitions are what let a consumer compute one, or see that it cannot.
        numbering = self.structure["numbering"]
        self.assertIn("7", numbering)
        level = numbering["7"]["levels"].get("0")
        self.assertIsNotNone(level, "level 0 of the list in use must be described")
        for field in ("start", "numFmt", "lvlText", "lvlRestart"):
            self.assertIn(field, level)

    def test_it_prints_no_number_of_its_own(self):
        # The division of labour: we hand over facts, the consumer computes and labels. This
        # checks the shape rather than the vocabulary — the note is allowed to talk about
        # clauses, and an earlier substring scan failed the moment it did.
        for entry in self.structure["paragraphs"]:
            self.assertEqual(set(entry) & {"clause", "number", "label", "text"}, set(),
                             "a computed number must not travel as if it were a fact")
        self.assertIn("derived", self.structure["note"])

    def test_a_document_with_no_numbering_gets_no_sidecar(self):
        import docx
        plain = os.path.join(self.root, "plain.docx")
        d = docx.Document()
        d.add_paragraph("Vienkāršs teksts bez numerācijas")
        d.save(plain)
        _md, _media, structure = normalize.docx_to_md(
            plain, os.path.join(self.root, "media"), "plain")
        self.assertIsNone(structure)


if __name__ == "__main__":
    unittest.main()
